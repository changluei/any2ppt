import json
from pathlib import Path

import pytest
from docx import Document

from app.ai.embeddings import HashEmbeddingProvider
from app.ai.evaluation import GoldenQuery, evaluate_retrieval
from app.ai.exceptions import IngestionError
from app.ai.ingestion import Chunk, parse_document, split_blocks
from app.ai.vector_store import ProjectVectorStore


FIXTURE = Path(__file__).parent / "fixtures" / "golden_retrieval.json"


def test_markdown_keeps_heading_and_line_range(tmp_path: Path):
    path = tmp_path / "lesson.md"
    path.write_text("# 蒸发\n\n水受热会加快蒸发。\n空气流动也会加快蒸发。", "utf-8")
    blocks = parse_document(path)
    assert blocks[0].heading == "蒸发"
    assert "第 3-4 行" in blocks[0].location
    assert blocks[0].content_hash


def test_docx_keeps_heading_paragraph_and_table_location(tmp_path: Path):
    path = tmp_path / "lesson.docx"
    document = Document()
    document.add_heading("水的变化", level=1)
    document.add_paragraph("水受热会蒸发。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "条件"
    table.rows[0].cells[1].text = "现象"
    document.save(path)
    blocks = parse_document(path)
    assert any(block.heading == "水的变化" and "段落" in block.location for block in blocks)
    assert any("表格 1" in block.location and "条件" in block.content for block in blocks)


def test_scanned_pdf_returns_clear_ocr_error(tmp_path: Path, monkeypatch):
    import pypdf

    class Page:
        @staticmethod
        def extract_text():
            return ""

    class Reader:
        pages = [Page(), Page()]

    monkeypatch.setattr(pypdf, "PdfReader", lambda _: Reader())
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"not-a-real-pdf")
    with pytest.raises(IngestionError, match="OCR"):
        parse_document(path)


def test_chunk_ids_are_stable_and_duplicate_content_is_removed():
    blocks = [
        type("Block", (), {"content": "水受热会蒸发。" * 8, "location": "第 1 页", "heading": "蒸发"})(),
        type("Block", (), {"content": "水受热会蒸发。" * 8, "location": "第 2 页", "heading": "蒸发"})(),
    ]
    first = split_blocks(blocks, "source-1", 50, 10)
    second = split_blocks(blocks, "source-1", 50, 10)
    assert [item.chunk_id for item in first] == [item.chunk_id for item in second]
    assert len({item.content_hash for item in first}) == len(first)


def _golden_store(tmp_path: Path):
    data = json.loads(FIXTURE.read_text("utf-8"))
    store = ProjectVectorStore(tmp_path, HashEmbeddingProvider(256), force_json=True)
    for row in data:
        if row["source_id"]:
            content_hash = f"hash-{row['source_id']}"
            store.add_documents(
                "golden-project",
                row["source_id"],
                row["filename"],
                [Chunk(f"chunk-{row['source_id']}", row["content"], row["category"], content_hash, row["category"])],
            )
    return store, data


def test_twenty_golden_queries_reach_top3_target(tmp_path: Path):
    store, data = _golden_store(tmp_path)
    queries = [GoldenQuery("golden-project", row["query"], row["source_id"]) for row in data]
    report = evaluate_retrieval(store, queries, top_k=3, min_score=0.12)
    assert report["total"] == 20
    assert report["top3_hit_rate"] >= 0.85
    assert report["empty_retrievals"] >= 2


def test_json_store_is_idempotent_deletable_and_project_isolated(tmp_path: Path):
    store = ProjectVectorStore(tmp_path, force_json=True)
    chunk = Chunk("c1", "水受热会蒸发", "第 1 页", "hash-1", "蒸发")
    store.add_documents("project-a", "source-a", "a.md", [chunk])
    store.add_documents("project-a", "source-a", "a.md", [chunk])
    store.add_documents("project-b", "source-b", "b.md", [Chunk("c2", "植物需要阳光", "第 1 页", "hash-2")])
    assert store.count("project-a") == 1
    assert all(row["project_id"] == "project-a" for row in store.similarity_search("project-a", "水受热", min_score=0))
    store.delete_by_source("project-a", "source-a")
    assert store.count("project-a") == 0
    assert store.count("project-b") == 1


def test_real_chroma_persists_and_filters_source(tmp_path: Path):
    store = ProjectVectorStore(tmp_path, HashEmbeddingProvider(128))
    store.add_documents("project/chroma", "source-a", "a.md", [Chunk("ca", "水受热会蒸发", "第 1 页", "ha")])
    store.add_documents("project/chroma", "source-b", "b.md", [Chunk("cb", "植物需要阳光", "第 1 页", "hb")])
    rows = store.similarity_search("project/chroma", "水受热蒸发", top_k=3, source_ids=["source-a"], min_score=0)
    assert rows and rows[0]["source_id"] == "source-a"
    assert ProjectVectorStore(tmp_path, HashEmbeddingProvider(128)).count("project/chroma") == 2
