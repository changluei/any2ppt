"""上传资料的解析、清洗与重叠切片。

支持 PDF、DOCX、Markdown 和纯文本。每个 Chunk 保留页码/标题/段落定位与
内容哈希，既用于引用追溯，也用于去重；此模块只生成 chunk，不直接写数据库
或向量库。
"""

from __future__ import annotations

import hashlib
import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from .exceptions import IngestionError


@dataclass(frozen=True)
class ParsedBlock:
    """从原文抽取的最小结构块及其可读位置。"""
    content: str
    location: str
    heading: str = ""
    page_number: int | None = None
    paragraph_index: int | None = None
    line_start: int | None = None
    line_end: int | None = None
    content_hash: str = ""


@dataclass(frozen=True)
class Chunk:
    """适合 embedding 的文本片段和可追溯元数据。"""
    chunk_id: str
    content: str
    location: str
    content_hash: str = ""
    heading: str = ""


def _clean(text: str) -> str:
    """规范换行和空白，降低格式噪声对检索的影响。"""
    lines = [re.sub(r"[ \t\u3000]+", " ", line).strip() for line in text.replace("\r", "").split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _hash(text: str) -> str:
    """生成稳定短哈希，作为 chunk 标识的一部分。"""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _text_blocks(text: str, *, markdown: bool) -> list[ParsedBlock]:
    """按 Markdown 标题或自然段切出文本结构块。"""
    lines = text.replace("\r", "").split("\n")
    blocks: list[ParsedBlock] = []
    heading = "正文"
    buffer: list[str] = []
    start_line = 1

    def flush(end_line: int) -> None:
        nonlocal buffer, start_line
        content = _clean("\n".join(buffer))
        if content:
            location = f"{heading} / 第 {start_line}-{end_line} 行" if markdown else f"第 {start_line}-{end_line} 行"
            blocks.append(
                ParsedBlock(
                    content=content,
                    location=location,
                    heading=heading if markdown else "",
                    line_start=start_line,
                    line_end=end_line,
                    content_hash=_hash(content),
                )
            )
        buffer = []

    for index, raw in enumerate(lines, 1):
        stripped = raw.strip()
        match = re.match(r"^#{1,6}\s+(.+)$", stripped) if markdown else None
        if match:
            flush(index - 1)
            heading = match.group(1).strip()
            start_line = index + 1
            continue
        if not stripped:
            flush(index - 1)
            start_line = index + 1
            continue
        if not buffer:
            start_line = index
        buffer.append(raw)
    flush(len(lines))
    return blocks


def _repeated_pdf_edges(page_texts: list[str]) -> set[str]:
    """识别多页重复页眉页脚，避免它们污染高频检索结果。"""
    if len(page_texts) < 3:
        return set()
    candidates: list[str] = []
    for text in page_texts:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        candidates.extend(lines[:1] + lines[-1:])
    counts = Counter(candidates)
    threshold = max(3, int(len(page_texts) * 0.6))
    return {line for line, count in counts.items() if count >= threshold and len(line) <= 120}


def parse_document(path: Path) -> list[ParsedBlock]:
    """按扩展名解析文件；无法提取有效文本时抛出 IngestionError。"""
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md"}:
            return _text_blocks(path.read_text(encoding="utf-8-sig"), markdown=suffix == ".md")

        if suffix == ".pdf":
            from pypdf import PdfReader

            page_texts = [(page.extract_text() or "").strip() for page in PdfReader(str(path)).pages]
            if not any(page_texts):
                raise IngestionError("扫描版 PDF 无可提取文本，需要 OCR，当前版本不支持")
            repeated = _repeated_pdf_edges(page_texts)
            blocks: list[ParsedBlock] = []
            for index, raw in enumerate(page_texts, 1):
                lines = [line for line in raw.splitlines() if line.strip() not in repeated]
                content = _clean("\n".join(lines))
                if content:
                    blocks.append(
                        ParsedBlock(
                            content=content,
                            location=f"第 {index} 页",
                            page_number=index,
                            content_hash=_hash(content),
                        )
                    )
            return blocks

        if suffix == ".docx":
            from docx import Document

            document = Document(str(path))
            result: list[ParsedBlock] = []
            heading = "正文"
            for index, paragraph in enumerate(document.paragraphs, 1):
                text = _clean(paragraph.text)
                if not text:
                    continue
                style_name = paragraph.style.name if paragraph.style else ""
                if style_name.startswith("Heading") or style_name.startswith("标题"):
                    heading = text
                    continue
                result.append(
                    ParsedBlock(
                        content=text,
                        location=f"{heading} / 段落 {index}",
                        heading=heading,
                        paragraph_index=index,
                        content_hash=_hash(text),
                    )
                )
            for table_index, table in enumerate(document.tables, 1):
                for row_index, row in enumerate(table.rows, 1):
                    text = _clean(" | ".join(cell.text for cell in row.cells))
                    if text:
                        result.append(
                            ParsedBlock(
                                content=text,
                                location=f"表格 {table_index} / 第 {row_index} 行",
                                heading=f"表格 {table_index}",
                                content_hash=_hash(text),
                            )
                        )
            return result
    except IngestionError:
        raise
    except (OSError, ValueError) as exc:
        raise IngestionError(f"资料解析失败：{path.name}") from exc
    raise IngestionError("仅支持 PDF、DOCX、TXT、Markdown")


def split_blocks(
    blocks: list[ParsedBlock],
    source_id: str,
    chunk_size: int = 500,
    overlap: int = 60,
) -> list[Chunk]:
    """按字符窗口切片，并保留 overlap 以减少语义在边界处断裂。"""
    if chunk_size < 10:
        raise ValueError("chunk_size 不能小于 10")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap 必须大于等于 0 且小于 chunk_size")
    chunks: list[Chunk] = []
    seen: set[str] = set()
    for block in blocks:
        start = 0
        while start < len(block.content):
            content = block.content[start : start + chunk_size].strip()
            if content:
                content_hash = _hash(content)
                if content_hash not in seen:
                    stable = f"{source_id}:{block.location}:{start}:{content_hash}"
                    chunk_id = hashlib.sha256(stable.encode("utf-8")).hexdigest()[:24]
                    location = block.location if len(block.content) <= chunk_size else f"{block.location} / 字符 {start + 1}-{start + len(content)}"
                    chunks.append(Chunk(chunk_id, content, location, content_hash, block.heading))
                    seen.add(content_hash)
            if start + chunk_size >= len(block.content):
                break
            start += chunk_size - overlap
    if not chunks:
        raise IngestionError("资料没有可索引的文本")
    return chunks
